#!/usr/bin/env python3
"""
Word文档格式诊断工具
用于排查不同电脑上docx文件解析失败的问题
"""

import sys
from pathlib import Path
import zipfile

def diagnose_docx(file_path: str):
    """诊断docx文件格式"""
    
    print("=" * 80)
    print("📋 Word文档格式诊断工具")
    print("=" * 80)
    print()
    
    file_path = Path(file_path)
    
    # 1. 检查文件是否存在
    print("1️⃣ 文件基本信息:")
    print(f"   路径: {file_path}")
    print(f"   存在: {file_path.exists()}")
    if not file_path.exists():
        print("   ❌ 文件不存在！")
        return
    
    print(f"   大小: {file_path.stat().st_size / 1024:.2f} KB")
    print(f"   扩展名: {file_path.suffix.lower()}")
    print()
    
    # 2. 检查文件扩展名
    print("2️⃣ 扩展名检查:")
    extension = file_path.suffix.lower()
    if extension not in ['.doc', '.docx']:
        print(f"   ⚠️  非标准Word扩展名: {extension}")
    else:
        print(f"   ✅ 标准Word扩展名: {extension}")
    print()
    
    # 3. 检查是否为有效的ZIP文件（docx本质是zip）
    print("3️⃣ ZIP结构检查:")
    try:
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            file_list = zip_ref.namelist()
            print(f"   ✅ 有效的ZIP文件")
            print(f"   文件数量: {len(file_list)}")
            
            # 检查关键文件
            critical_files = [
                'word/document.xml',
                '[Content_Types].xml',
                '_rels/.rels'
            ]
            
            print("\n   关键文件检查:")
            for cf in critical_files:
                if cf in file_list:
                    print(f"   ✅ {cf}")
                else:
                    print(f"   ❌ {cf} (缺失)")
            
            # 显示前10个文件
            print(f"\n   前10个文件:")
            for f in file_list[:10]:
                print(f"      - {f}")
                
    except zipfile.BadZipFile:
        print(f"   ❌ 不是有效的ZIP文件（可能是.doc格式）")
        print(f"   提示: .doc是旧格式，需要转换为.docx")
    except Exception as e:
        print(f"   ❌ ZIP检查失败: {e}")
    print()
    
    # 4. 尝试用python-docx读取
    print("4️⃣ python-docx 兼容性:")
    try:
        from docx import Document
        doc = Document(file_path)
        print(f"   ✅ python-docx 可以读取")
        print(f"   段落数: {len(doc.paragraphs)}")
        print(f"   表格数: {len(doc.tables)}")
    except Exception as e:
        print(f"   ❌ python-docx 读取失败: {e}")
    print()
    
    # 5. 尝试用docxtpl读取
    print("5️⃣ docxtpl (DocxTemplate) 兼容性:")
    try:
        from docxtpl import DocxTemplate
        doc = DocxTemplate(file_path)
        vars = doc.undeclared_template_variables
        print(f"   ✅ docxtpl 可以读取")
        print(f"   模板变量数: {len(vars)}")
        if vars:
            print(f"   变量: {list(vars)[:5]}")
    except Exception as e:
        print(f"   ❌ docxtpl 读取失败: {e}")
        print(f"   错误类型: {type(e).__name__}")
    print()
    
    # 6. 检查文件编码和特殊字符
    print("6️⃣ 文件名编码检查:")
    try:
        # 检查文件名是否包含特殊字符
        filename = file_path.name
        print(f"   文件名: {filename}")
        print(f"   文件名编码: {filename.encode('utf-8')}")
        
        # 检查是否有非ASCII字符
        if all(ord(c) < 128 for c in filename):
            print(f"   ✅ 纯ASCII文件名")
        else:
            print(f"   ⚠️  包含非ASCII字符（可能在某些系统上有问题）")
    except Exception as e:
        print(f"   ❌ 编码检查失败: {e}")
    print()
    
    # 7. 总结建议
    print("=" * 80)
    print("📝 诊断总结和建议:")
    print("=" * 80)
    
    suggestions = []
    
    if extension == '.doc':
        suggestions.append("• 文件是旧的.doc格式，请转换为.docx格式")
    
    try:
        with zipfile.ZipFile(file_path, 'r'):
            pass
    except:
        suggestions.append("• 文件不是有效的.docx格式（非ZIP结构）")
        suggestions.append("• 建议用Microsoft Word或WPS重新保存为.docx")
    
    if not suggestions:
        suggestions.append("✅ 文件格式检查通过，可能是其他环境问题")
        suggestions.append("• 检查python-docx和docxtpl库版本是否一致")
        suggestions.append("• 检查Python版本是否一致")
        suggestions.append("• 尝试在另一台电脑上重新保存文件")
    
    for s in suggestions:
        print(s)
    
    print()
    print("=" * 80)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python diagnose_docx.py <docx文件路径>")
        print()
        print("示例:")
        print("  python diagnose_docx.py template.docx")
        print("  python diagnose_docx.py /path/to/your/file.docx")
        sys.exit(1)
    
    diagnose_docx(sys.argv[1])

