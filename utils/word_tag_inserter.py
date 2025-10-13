"""
Word文档标签插入工具
用于在现有Word文档中插入智能标签
"""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Tuple
import os
from pathlib import Path


class WordTagInserter:
    """Word文档标签插入器"""
    
    def __init__(self):
        self.supported_tags = self._init_supported_tags()
    
    def _init_supported_tags(self) -> List[Dict[str, str]]:
        """初始化支持的标签列表"""
        return [
            # 基本信息
            {'tag': 'course_name', 'desc': '课程名称', 'category': '基本信息'},
            {'tag': 'teacher_name', 'desc': '授课教师', 'category': '基本信息'},
            {'tag': 'class_name', 'desc': '授课班级', 'category': '基本信息'},
            {'tag': 'lesson_number', 'desc': '课次', 'category': '基本信息'},
            {'tag': 'lesson_title', 'desc': '课题', 'category': '基本信息'},
            {'tag': 'teaching_hours', 'desc': '学时', 'category': '基本信息'},
            {'tag': 'chapter_section', 'desc': '授课章节', 'category': '基本信息'},
            
            # 教学目标
            {'tag': 'ideological_goals', 'desc': '思政育人目标', 'category': '教学目标'},
            {'tag': 'knowledge_goals', 'desc': '知识目标', 'category': '教学目标'},
            {'tag': 'ability_goals', 'desc': '能力目标', 'category': '教学目标'},
            {'tag': 'ideological_elements', 'desc': '思政元素', 'category': '教学目标'},
            
            # 教学重难点
            {'tag': 'teaching_focus', 'desc': '教学重点', 'category': '教学重难点'},
            {'tag': 'focus_solutions', 'desc': '重点解决措施', 'category': '教学重难点'},
            {'tag': 'teaching_difficulty', 'desc': '教学难点', 'category': '教学重难点'},
            {'tag': 'difficulty_solutions', 'desc': '难点解决措施', 'category': '教学重难点'},
            
            # 教学方法
            {'tag': 'teaching_methods', 'desc': '教法', 'category': '教学方法'},
            {'tag': 'learning_methods', 'desc': '学法', 'category': '教学方法'},
            {'tag': 'teaching_resources', 'desc': '教学资源', 'category': '教学方法'},
            
            # 课前预习
            {'tag': 'preview_content', 'desc': '课前预习-教学内容', 'category': '课前预习'},
            {'tag': 'preview_teacher', 'desc': '课前预习-教师活动', 'category': '课前预习'},
            {'tag': 'preview_student', 'desc': '课前预习-学生活动', 'category': '课前预习'},
            {'tag': 'preview_intention', 'desc': '课前预习-设计意图', 'category': '课前预习'},
            
            # 自主学习
            {'tag': 'self_learning_content', 'desc': '自主学习-预习内容', 'category': '自主学习'},
            {'tag': 'self_learning_teacher', 'desc': '自主学习-教师活动', 'category': '自主学习'},
            {'tag': 'self_learning_student', 'desc': '自主学习-学生活动', 'category': '自主学习'},
            {'tag': 'self_learning_intention', 'desc': '自主学习-设计意图', 'category': '自主学习'},
            
            # 新课导入
            {'tag': 'introduction_content', 'desc': '新课导入-教学内容', 'category': '新课导入'},
            {'tag': 'introduction_teacher', 'desc': '新课导入-教师活动', 'category': '新课导入'},
            {'tag': 'introduction_student', 'desc': '新课导入-学生活动', 'category': '新课导入'},
            {'tag': 'introduction_intention', 'desc': '新课导入-设计意图', 'category': '新课导入'},
            
            # 预习反馈
            {'tag': 'feedback_content', 'desc': '预习反馈-教学内容', 'category': '预习反馈'},
            {'tag': 'feedback_teacher', 'desc': '预习反馈-教师活动', 'category': '预习反馈'},
            {'tag': 'feedback_student', 'desc': '预习反馈-学生活动', 'category': '预习反馈'},
            {'tag': 'feedback_intention', 'desc': '预习反馈-设计意图', 'category': '预习反馈'},
            
            # 新课讲授
            {'tag': 'teaching_content', 'desc': '新课讲授-教学内容', 'category': '新课讲授'},
            {'tag': 'teaching_teacher', 'desc': '新课讲授-教师活动', 'category': '新课讲授'},
            {'tag': 'teaching_student', 'desc': '新课讲授-学生活动', 'category': '新课讲授'},
            {'tag': 'teaching_intention', 'desc': '新课讲授-设计意图', 'category': '新课讲授'},
            
            # 实践环节
            {'tag': 'practice_content', 'desc': '实践-教学内容', 'category': '实践环节'},
            {'tag': 'practice_teacher', 'desc': '实践-教师活动', 'category': '实践环节'},
            {'tag': 'practice_student', 'desc': '实践-学生活动', 'category': '实践环节'},
            {'tag': 'practice_intention', 'desc': '实践-设计意图', 'category': '实践环节'},
            
            # 展示环节
            {'tag': 'presentation_content', 'desc': '展示-教学内容', 'category': '展示环节'},
            {'tag': 'presentation_teacher', 'desc': '展示-教师活动', 'category': '展示环节'},
            {'tag': 'presentation_student', 'desc': '展示-学生活动', 'category': '展示环节'},
            {'tag': 'presentation_intention', 'desc': '展示-设计意图', 'category': '展示环节'},
            
            # 评价环节
            {'tag': 'evaluation_content', 'desc': '评价-教学内容', 'category': '评价环节'},
            {'tag': 'evaluation_teacher', 'desc': '评价-教师活动', 'category': '评价环节'},
            {'tag': 'evaluation_student', 'desc': '评价-学生活动', 'category': '评价环节'},
            {'tag': 'evaluation_intention', 'desc': '评价-设计意图', 'category': '评价环节'},
            
            # 课后作业
            {'tag': 'homework_content', 'desc': '课后作业-教学内容', 'category': '课后作业'},
            {'tag': 'homework_teacher', 'desc': '课后作业-教师活动', 'category': '课后作业'},
            {'tag': 'homework_student', 'desc': '课后作业-学生活动', 'category': '课后作业'},
            {'tag': 'homework_intention', 'desc': '课后作业-设计意图', 'category': '课后作业'},
            
            # 阅读延伸
            {'tag': 'extension_content', 'desc': '阅读延伸-教学内容', 'category': '阅读延伸'},
            {'tag': 'extension_teacher', 'desc': '阅读延伸-教师活动', 'category': '阅读延伸'},
            {'tag': 'extension_student', 'desc': '阅读延伸-学生活动', 'category': '阅读延伸'},
            {'tag': 'extension_intention', 'desc': '阅读延伸-设计意图', 'category': '阅读延伸'},
            
            # 教学反思
            {'tag': 'reflection_effects', 'desc': '教学反思-目标效果', 'category': '教学反思'},
            {'tag': 'reflection_improvements', 'desc': '教学反思-反思改进', 'category': '教学反思'},
        ]
    
    def extract_document_structure(self, docx_path: str) -> Dict:
        """
        提取Word文档完整结构（用于可视化显示）
        
        Returns:
            包含完整文档布局的字典，按照文档顺序排列段落和表格
        """
        try:
            doc = Document(docx_path)
            structure = {
                'elements': [],  # 按顺序存储所有元素（段落和表格）
                'total_paragraphs': 0,
                'total_tables': 0
            }
            
            # 获取文档中所有元素的顺序
            para_index = 0
            table_index = 0
            
            # 遍历文档的所有元素（段落和表格混合）
            for element in doc.element.body:
                # 检查是否是段落
                if element.tag.endswith('p'):
                    try:
                        para = doc.paragraphs[para_index]
                        structure['elements'].append({
                            'type': 'paragraph',
                            'index': para_index,
                            'text': para.text,
                            'style': para.style.name if para.style else 'Normal',
                            'alignment': str(para.alignment) if para.alignment else 'LEFT'
                        })
                        para_index += 1
                    except:
                        pass
                
                # 检查是否是表格
                elif element.tag.endswith('tbl'):
                    try:
                        table = doc.tables[table_index]
                        table_data = {
                            'type': 'table',
                            'index': table_index,
                            'rows': len(table.rows),
                            'cols': len(table.columns) if table.rows else 0,
                            'cells': [],
                            'cell_map': []  # 🔥 新增：单元格映射表
                        }
                        
                        # 提取所有单元格，并标记合并单元格
                        for row_idx, row in enumerate(table.rows):
                            row_cells = []
                            row_map = []  # 该行的单元格映射
                            
                            seen_cell_ids = {}  # 记录已见过的单元格ID
                            
                            for col_idx, cell in enumerate(row.cells):
                                cell_text = cell.text.strip()
                                cell_id = id(cell)
                                
                                # 检查是否是合并单元格（与之前的单元格ID相同）
                                is_merged = cell_id in seen_cell_ids
                                first_col = seen_cell_ids.get(cell_id, col_idx)
                                
                                row_cells.append({
                                    'row': row_idx,
                                    'col': col_idx,
                                    'text': cell_text,
                                    'is_empty': not bool(cell_text),
                                    'is_merged': is_merged,
                                    'first_col': first_col  # 合并单元格的第一列索引
                                })
                                
                                # 只为独立单元格创建映射
                                if not is_merged:
                                    row_map.append({
                                        'col': col_idx,
                                        'text': cell_text[:20],
                                        'is_empty': not bool(cell_text)
                                    })
                                    seen_cell_ids[cell_id] = col_idx
                            
                            table_data['cells'].append(row_cells)
                            table_data['cell_map'].append(row_map)
                        
                        structure['elements'].append(table_data)
                        table_index += 1
                    except:
                        pass
            
            structure['total_paragraphs'] = para_index
            structure['total_tables'] = table_index
            
            return structure
            
        except Exception as e:
            print(f"❌ 提取文档结构失败: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def insert_tag_to_document(self, docx_path: str, location: Dict, tag_name: str, output_path: str = None) -> Tuple[str, bool]:
        """
        在指定位置插入标签
        
        Args:
            docx_path: Word文档路径
            location: 位置信息 {'type': 'paragraph'/'table', 'index': int, 'row': int, 'col': int}
            tag_name: 标签名
            output_path: 输出路径，如不指定则覆盖原文件
            
        Returns:
            (output_path, success)
        """
        try:
            doc = Document(docx_path)
            tag_text = f"{{{{{tag_name}}}}}"  # 生成 {{tag_name}}
            
            print("=" * 80)
            print(f"📌 插入标签调试信息")
            print("=" * 80)
            print(f"📄 文档信息:")
            print(f"   总段落数: {len(doc.paragraphs)}")
            print(f"   总表格数: {len(doc.tables)}")
            print(f"🏷️  标签名: {tag_name}")
            print(f"📍 位置信息: {location}")
            
            if location['type'] == 'paragraph':
                # 🔥 优先使用文本定位，其次使用索引
                para_text = location.get('text', '')
                index = location['index']
                
                para = None
                
                if para_text:
                    # 方法1: 根据文本查找段落
                    print(f"🎯 尝试根据文本查找段落: \"{para_text[:30]}...\"")
                    for i, p in enumerate(doc.paragraphs):
                        # 去除空格后比较（因为Word可能有额外空格）
                        if p.text.strip().startswith(para_text.strip()[:20]):
                            para = p
                            index = i
                            print(f"   ✅ 找到匹配段落 {i}: \"{p.text[:50]}...\"")
                            break
                    
                    if not para:
                        print(f"   ⚠️ 未找到匹配文本的段落，尝试使用索引 {index}")
                
                # 方法2: 使用索引（如果文本查找失败）
                if not para:
                    print(f"🎯 尝试插入到段落 {index}")
                    
                    if index >= len(doc.paragraphs):
                        print(f"⚠️ 段落索引 {index} 超出范围，最大索引: {len(doc.paragraphs) - 1}")
                        return None, False
                    
                    para = doc.paragraphs[index]
                    print(f"   段落内容: {para.text[:50]}...")
                
                # 在段落末尾插入
                run = para.add_run(f" {tag_text}")
                # 设置标签样式
                run.font.color.rgb = RGBColor(16, 163, 127)
                run.font.size = Pt(10.5)
                
            elif location['type'] == 'table':
                table_idx = location['index']
                row_idx = location['row']
                col_idx = location['col']
                
                print(f"🎯 尝试插入到表格 {table_idx}, 行 {row_idx}, 列 {col_idx}")
                
                if table_idx >= len(doc.tables):
                    print(f"⚠️ 表格索引 {table_idx} 超出范围，最大索引: {len(doc.tables) - 1}")
                    return None, False
                
                table = doc.tables[table_idx]
                print(f"   表格总行数: {len(table.rows)}")
                print(f"   表格总列数: {len(table.columns)}")
                
                if row_idx >= len(table.rows):
                    print(f"⚠️ 行索引 {row_idx} 超出范围，最大索引: {len(table.rows) - 1}")
                    return None, False
                
                row = table.rows[row_idx]
                print(f"   该行的单元格数量: {len(row.cells)}")
                
                # 🔥 调试：显示该行所有单元格的文本
                print(f"   该行所有单元格内容:")
                for i, c in enumerate(row.cells):
                    preview = c.text.strip()[:30] if c.text.strip() else "(空)"
                    print(f"      [{row_idx},{i}]: \"{preview}\"")
                
                if col_idx >= len(row.cells):
                    print(f"⚠️ 列索引 {col_idx} 超出范围，最大索引: {len(row.cells) - 1}")
                    return None, False
                
                cell = row.cells[col_idx]
                cell_text = cell.text[:50] if cell.text else "(空)"
                cell_text_stripped = cell.text.strip() if cell.text else ""
                
                print(f"   🎯 目标单元格 [{row_idx},{col_idx}] 当前内容: \"{cell_text}\"")
                
                # 🔥 智能提示：如果单元格已有非标签内容，给出警告
                if cell_text_stripped and not cell_text_stripped.startswith('{{'):
                    print(f"   ⚠️  注意：该单元格已有内容，标签将追加在内容后")
                    print(f"   💡 提示：如果要插入到空白单元格，请点击更右边的单元格")
                
                # 在cell的第一个段落末尾追加标签（不管是否有内容）
                if cell.paragraphs:
                    # 在第一个段落的末尾追加
                    para = cell.paragraphs[0]
                    run = para.add_run(f" {tag_text}")
                    run.font.color.rgb = RGBColor(16, 163, 127)
                    run.font.size = Pt(10.5)
                    print(f"   ✅ 标签追加在原内容后")
                else:
                    # 如果没有段落，创建新内容
                    cell.text = tag_text
                    print(f"   ✅ 创建新内容")
            
            # 保存
            if output_path is None:
                output_path = docx_path
            
            doc.save(output_path)
            print(f"✅ 标签 {tag_text} 已成功插入到文件")
            print("=" * 80)
            return output_path, True
            
        except Exception as e:
            print(f"❌ 插入标签失败: {e}")
            import traceback
            traceback.print_exc()
            print("=" * 80)
            return None, False
    
    def get_tags_by_category(self) -> Dict[str, List[Dict]]:
        """获取按分类组织的标签列表"""
        categories = {}
        for tag_info in self.supported_tags:
            category = tag_info['category']
            if category not in categories:
                categories[category] = []
            categories[category].append(tag_info)
        return categories


if __name__ == "__main__":
    # 测试代码
    inserter = WordTagInserter()
    
    print("✅ WordTagInserter 初始化成功")
    print(f"📋 支持 {len(inserter.supported_tags)} 个标签")
    
    categories = inserter.get_tags_by_category()
    print(f"\n📂 标签分类({len(categories)}个):")
    for cat, tags in categories.items():
        print(f"  - {cat}: {len(tags)}个")

