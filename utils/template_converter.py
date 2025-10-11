"""Template document conversion utilities - converts DOC/DOCX to images"""

import tempfile
import os
from pathlib import Path
from typing import List

# Import handling for optional dependencies
try:
    import aspose.words as aw
    ASPOSE_AVAILABLE = True
except ImportError:
    ASPOSE_AVAILABLE = False
    
try:
    from docx2python import docx2python
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    DOC_CONVERSION_AVAILABLE = True
except ImportError:
    DOC_CONVERSION_AVAILABLE = False


class TemplateConverter:
    """Template document converter for handling DOC/DOCX to image conversion"""
    
    @staticmethod
    def convert_to_images(file_path: str) -> List[str]:
        """
        Convert DOC/DOCX template files to images for parsing
        
        Args:
            file_path: Path to the template file
            
        Returns:
            List of image file paths
        """
        try:
            if ASPOSE_AVAILABLE:
                return TemplateConverter._convert_with_aspose(file_path)
            elif DOC_CONVERSION_AVAILABLE:
                return TemplateConverter._convert_with_docx2python(file_path)
            else:
                return TemplateConverter._convert_with_simple_render(file_path)
        except Exception as e:
            print(f"Template conversion error: {e}")
            return []
    
    @staticmethod
    def _convert_with_aspose(file_path: str) -> List[str]:
        """Convert using Aspose.Words (best quality)"""
        try:
            # Load document
            doc = aw.Document(file_path)
            
            # Create save options for high-quality images
            save_options = aw.saving.ImageSaveOptions(aw.SaveFormat.PNG)
            save_options.horizontal_resolution = 300
            save_options.vertical_resolution = 300
            save_options.scale = 1.0
            
            # Convert each page
            image_paths = []
            temp_dir = tempfile.mkdtemp()
            
            for page_index in range(doc.page_count):
                save_options.page_set = aw.saving.PageSet([page_index])
                output_path = os.path.join(temp_dir, f"template_page_{page_index + 1}.png")
                doc.save(output_path, save_options)
                image_paths.append(output_path)
            
            print(f"✅ Aspose conversion successful: {len(image_paths)} pages")
            return image_paths
        except Exception as e:
            print(f"❌ Aspose conversion failed: {e}")
            return []
    
    @staticmethod
    def _convert_with_docx2python(file_path: str) -> List[str]:
        """Convert using docx2python (medium quality)"""
        try:
            content = docx2python(file_path)
            
            image_paths = []
            temp_dir = tempfile.mkdtemp()
            
            # Render document content as image
            fig, ax = plt.subplots(figsize=(8.5, 11))  # A4 paper size
            ax.axis('off')
            
            # Process text content
            text_content = content.text
            if len(text_content) > 2000:  # Limit length to avoid display issues
                text_content = text_content[:2000] + "\n\n[内容过长，已截断...]"
            
            # Add text to image with better formatting
            ax.text(0.05, 0.95, text_content, 
                   transform=ax.transAxes, 
                   fontsize=10, 
                   verticalalignment='top',
                   wrap=True,
                   family='monospace')
            
            # Save image
            output_path = os.path.join(temp_dir, "template_page_1.png")
            plt.tight_layout()
            plt.savefig(output_path, dpi=300, bbox_inches='tight', 
                       facecolor='white', edgecolor='none')
            plt.close()
            
            image_paths.append(output_path)
            print(f"✅ docx2python conversion successful: 1 page")
            return image_paths
        except Exception as e:
            print(f"❌ docx2python conversion failed: {e}")
            return []
    
    @staticmethod
    def _convert_with_simple_render(file_path: str) -> List[str]:
        """Simple text rendering method (fallback)"""
        try:
            text_content = ""
            file_extension = Path(file_path).suffix.lower()
            
            # Try to read docx file content
            if file_extension == '.docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    if not text_content.strip():
                        text_content = "文档内容为空或无法提取文本"
                except ImportError:
                    text_content = """
⚠️ 无法完整读取DOCX文件

缺少必要的库，建议：
1. 安装依赖: pip install python-docx
2. 或将文档另存为图片格式上传
3. 或安装高级转换库: pip install aspose-words

系统将使用默认模板结构。
"""
            else:
                # For DOC files, provide instruction
                text_content = """
⚠️ 检测到DOC格式文件

由于缺少转换库，无法直接处理DOC文件。

建议方案：
1. 将DOC文件另存为DOCX格式后重新上传
2. 或将文档截图保存为图片上传
3. 或安装转换库: pip install aspose-words

系统将使用默认模板结构继续。
"""
            
            # Create simple text image using PIL
            from PIL import Image, ImageDraw, ImageFont
            
            temp_dir = tempfile.mkdtemp()
            
            # Create blank image (A4 proportions)
            img_width, img_height = 850, 1100
            img = Image.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(img)
            
            # Try to use better font
            try:
                font = ImageFont.truetype("arial.ttf", 14)
                title_font = ImageFont.truetype("arial.ttf", 18)
            except:
                try:
                    # Try common Chinese fonts
                    font = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 14)
                    title_font = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 18)
                except:
                    # Use default font
                    font = ImageFont.load_default()
                    title_font = ImageFont.load_default()
            
            # Add title
            draw.text((30, 30), "教案模板预览", font=title_font, fill='black')
            
            # Text wrapping and rendering
            lines = []
            for line in text_content.split('\n'):
                if not line.strip():
                    lines.append('')
                    continue
                # Wrap long lines
                if len(line) > 70:
                    while len(line) > 70:
                        lines.append(line[:70])
                        line = line[70:]
                    if line:
                        lines.append(line)
                else:
                    lines.append(line)
            
            # Draw text
            y_position = 80
            line_height = 25
            max_lines = 40
            
            for i, line in enumerate(lines[:max_lines]):
                try:
                    draw.text((30, y_position), line, font=font, fill='black')
                    y_position += line_height
                except:
                    # Handle encoding issues
                    draw.text((30, y_position), "[无法显示的字符]", font=font, fill='gray')
                    y_position += line_height
            
            if len(lines) > max_lines:
                draw.text((30, y_position), f"... (还有 {len(lines) - max_lines} 行未显示)", 
                         font=font, fill='gray')
            
            # Save image
            output_path = os.path.join(temp_dir, "template_simple.png")
            img.save(output_path, quality=95)
            
            print(f"⚠️ 使用简单渲染模式: 1 page")
            return [output_path]
            
        except Exception as e:
            print(f"❌ Simple rendering failed: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    @staticmethod
    def is_supported_format(file_path: str) -> bool:
        """Check if file format is supported"""
        supported_extensions = ['.doc', '.docx', '.jpg', '.jpeg', '.png', '.bmp', '.gif']
        file_extension = Path(file_path).suffix.lower()
        return file_extension in supported_extensions
    
    @staticmethod
    def get_conversion_method() -> str:
        """Get the available conversion method"""
        if ASPOSE_AVAILABLE:
            return "Aspose.Words (高质量)"
        elif DOC_CONVERSION_AVAILABLE:
            return "docx2python (中等质量)"
        else:
            return "简单渲染 (基础质量)"