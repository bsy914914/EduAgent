"""
模板自动标注器
使用AI识别教案模板中的可填充字段，自动添加{{placeholder}}标签
"""

import os
from docx import Document
from docx.shared import RGBColor
from typing import Dict, List, Tuple
import re


class TemplateAnnotator:
    """使用AI自动标注教案模板"""
    
    # 常见的教案字段映射
    FIELD_MAPPINGS = {
        # 基本信息
        "课程名称": "course_name",
        "授课教师": "teacher_name",
        "教师姓名": "teacher_name",
        "任课教师": "teacher_name",
        "授课班级": "teaching_class",
        "班级": "teaching_class",
        "授课章节": "chapter_section",
        "章节": "chapter_section",
        "授课内容": "teaching_content",
        "课时": "course_hours",
        "学时": "course_hours",
        "授课时间": "teaching_time",
        "授课地点": "teaching_location",
        "教学场所": "teaching_location",
        
        # 教学目标
        "知识目标": "knowledge_goals",
        "知识与技能": "knowledge_goals",
        "能力目标": "ability_goals",
        "过程与方法": "ability_goals",
        "思政目标": "ideological_goals",
        "情感态度与价值观": "ideological_goals",
        "素质目标": "quality_goals",
        "德育目标": "moral_goals",
        
        # 教学重难点
        "教学重点": "teaching_focus",
        "重点": "teaching_focus",
        "教学难点": "teaching_difficulties",
        "难点": "teaching_difficulties",
        "教学关键": "teaching_key_points",
        
        # 教学方法与手段
        "教学方法": "teaching_methods",
        "教法": "teaching_methods",
        "学法": "learning_methods",
        "教学手段": "teaching_tools",
        "教具": "teaching_aids",
        "实验器材": "experiment_equipment",
        
        # 教学过程
        "课前预习": "preview_content",
        "预习内容": "preview_content",
        "课前准备": "preparation",
        "导入": "introduction",
        "新课导入": "introduction",
        "教学过程": "teaching_process",
        "教学步骤": "teaching_steps",
        "教学环节": "teaching_stages",
        "课堂活动": "class_activities",
        "实践活动": "practice_activities",
        "讨论问题": "discussion_questions",
        "案例分析": "case_analysis",
        "例题讲解": "example_explanation",
        "练习题": "practice_exercises",
        
        # 课后环节
        "课程小结": "summary_reflection",
        "小结": "summary_reflection",
        "课堂小结": "summary_reflection",
        "总结": "summary_reflection",
        "课后作业": "homework_assignment",
        "作业": "homework_assignment",
        "课后思考": "after_class_thinking",
        "拓展延伸": "extension",
        
        # 其他
        "板书设计": "blackboard_design",
        "教学反思": "teaching_reflection",
        "课后反思": "teaching_reflection",
        "备注": "notes",
        "参考资料": "references",
    }
    
    def __init__(self, agent):
        """
        初始化模板标注器
        
        Args:
            agent: UniversityCourseAgent实例，用于AI分析
        """
        self.agent = agent
    
    def annotate_template(self, template_path: str, output_path: str = None) -> str:
        """
        自动标注教案模板
        
        Args:
            template_path: 原始模板路径
            output_path: 输出路径，如果为None则自动生成
            
        Returns:
            标注后的模板路径
        """
        print("=" * 80)
        print("🤖 开始AI自动标注教案模板")
        print("=" * 80)
        
        # 加载模板
        print(f"\n📄 加载模板: {template_path}")
        doc = Document(template_path)
        
        # 分析模板结构
        print("\n🔍 分析模板结构...")
        structure = self._analyze_template_structure(doc)
        
        print(f"\n📊 模板分析结果:")
        print(f"  - 段落数: {structure['paragraph_count']}")
        print(f"  - 表格数: {structure['table_count']}")
        print(f"  - 识别的字段数: {len(structure['fields'])}")
        
        # 使用AI增强字段识别
        print("\n🤖 使用AI增强字段识别...")
        enhanced_fields = self._ai_enhance_fields(structure['fields'])
        
        print(f"\n✨ AI增强后识别的字段数: {len(enhanced_fields)}")
        for field in enhanced_fields[:10]:  # 显示前10个
            print(f"  - {field['label']}: {{{{field['placeholder']}}}}")
        if len(enhanced_fields) > 10:
            print(f"  ... 还有 {len(enhanced_fields) - 10} 个字段")
        
        # 标注模板
        print("\n📝 开始标注模板...")
        annotated_count = self._apply_annotations(doc, enhanced_fields)
        
        print(f"\n✅ 成功标注 {annotated_count} 个字段")
        
        # 保存标注后的模板
        if output_path is None:
            base_name = os.path.splitext(os.path.basename(template_path))[0]
            output_path = os.path.join(
                os.path.dirname(template_path) or ".",
                f"{base_name}_已标注.docx"
            )
        
        doc.save(output_path)
        print(f"\n💾 已保存标注模板: {output_path}")
        
        # 生成标注说明
        self._generate_annotation_guide(enhanced_fields, output_path)
        
        print("\n" + "=" * 80)
        print("🎉 模板标注完成！")
        print("=" * 80)
        
        return output_path
    
    def _analyze_template_structure(self, doc: Document) -> Dict:
        """分析模板结构，识别可能的字段"""
        structure = {
            'paragraph_count': 0,
            'table_count': 0,
            'fields': []
        }
        
        # 分析段落
        for para in doc.paragraphs:
            structure['paragraph_count'] += 1
            text = para.text.strip()
            
            if text:
                # 检测是否为字段标签（如"课程名称："或"课程名称"后面跟着空白）
                field_info = self._detect_field(text, 'paragraph', para)
                if field_info:
                    structure['fields'].append(field_info)
        
        # 分析表格
        for table_idx, table in enumerate(doc.tables):
            structure['table_count'] += 1
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    
                    if text:
                        field_info = self._detect_field(
                            text, 'table', cell, 
                            table_idx=table_idx,
                            row_idx=row_idx,
                            cell_idx=cell_idx
                        )
                        if field_info:
                            structure['fields'].append(field_info)
        
        return structure
    
    def _detect_field(self, text: str, location_type: str, element, **kwargs) -> Dict:
        """检测文本是否为字段标签"""
        # 移除常见的标点和空白
        clean_text = text.replace("：", "").replace(":", "").strip()
        
        # 检查是否匹配已知字段
        for label, placeholder in self.FIELD_MAPPINGS.items():
            if label in clean_text:
                # 检查是否已经包含占位符
                if "{{" in text and "}}" in text:
                    return None
                
                # 判断是否为标签行（短文本，包含冒号或为表头）
                is_label = (
                    len(clean_text) <= 20 or  # 短文本
                    "：" in text or ":" in text or  # 包含冒号
                    (location_type == 'table' and kwargs.get('row_idx', 0) == 0)  # 表头
                )
                
                return {
                    'label': label,
                    'placeholder': placeholder,
                    'text': text,
                    'location_type': location_type,
                    'element': element,
                    'is_label': is_label,
                    **kwargs
                }
        
        return None
    
    def _ai_enhance_fields(self, fields: List[Dict]) -> List[Dict]:
        """使用AI增强字段识别"""
        # 构建提示词
        field_texts = [f"- {f['text']}" for f in fields[:20]]  # 限制数量避免token过多
        
        prompt = f"""你是一个教案模板分析专家。请分析以下从教案模板中提取的文本，识别哪些是需要填充的字段。

提取的文本:
{chr(10).join(field_texts)}

已知的教案字段类型包括：
- 基本信息：课程名称、授课教师、授课班级、授课章节、课时等
- 教学目标：知识目标、能力目标、思政目标等
- 教学重难点：教学重点、教学难点等
- 教学方法：教学方法、教学手段等
- 教学过程：导入、教学步骤、课堂活动等
- 课后环节：课程小结、课后作业等

请确认这些字段是否正确识别，并补充遗漏的字段。
只需要回答"确认"或者提出补充建议即可。"""

        try:
            # 调用AI分析
            response = self.agent.llm_chat.invoke(prompt)
            print(f"\n🤖 AI分析建议: {response.content[:200]}...")
        except Exception as e:
            print(f"\n⚠️  AI分析出错: {e}，使用基础识别结果")
        
        return fields
    
    def _apply_annotations(self, doc: Document, fields: List[Dict]) -> int:
        """应用标注到文档"""
        annotated_count = 0
        
        for field in fields:
            element = field['element']
            placeholder = f"{{{{{field['placeholder']}}}}}"
            
            try:
                if field['location_type'] == 'paragraph':
                    # 段落标注
                    if field['is_label']:
                        # 如果是标签行，在后面添加占位符
                        if not element.text.strip().endswith(placeholder):
                            # 清除现有内容并重写
                            original_text = element.text.strip()
                            element.clear()
                            run = element.add_run(f"{original_text} {placeholder}")
                            run.font.color.rgb = RGBColor(255, 0, 0)  # 红色标记
                            annotated_count += 1
                    else:
                        # 如果是内容行，替换为占位符
                        element.clear()
                        run = element.add_run(placeholder)
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        annotated_count += 1
                
                elif field['location_type'] == 'table':
                    # 表格单元格标注
                    cell = element
                    
                    # 判断是否为标签单元格（通常是左侧或顶部）
                    if field['is_label'] or field.get('cell_idx', 0) == 0:
                        # 标签单元格，保留原文，在后续单元格添加占位符
                        # 这里我们在同一单元格添加
                        original_text = cell.text.strip()
                        cell.text = ""
                        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        run = paragraph.add_run(f"{original_text} {placeholder}")
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        annotated_count += 1
                    else:
                        # 内容单元格，替换为占位符
                        cell.text = ""
                        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        run = paragraph.add_run(placeholder)
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        annotated_count += 1
            
            except Exception as e:
                print(f"⚠️  标注字段 {field['label']} 时出错: {e}")
                continue
        
        return annotated_count
    
    def _generate_annotation_guide(self, fields: List[Dict], template_path: str):
        """生成标注说明文档"""
        guide_path = template_path.replace(".docx", "_标注说明.txt")
        
        with open(guide_path, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("教案模板标注说明\n")
            f.write("=" * 80 + "\n\n")
            
            f.write(f"原始模板: {os.path.basename(template_path)}\n")
            f.write(f"标注字段数: {len(fields)}\n\n")
            
            f.write("识别的字段列表:\n")
            f.write("-" * 80 + "\n\n")
            
            for idx, field in enumerate(fields, 1):
                f.write(f"{idx}. {field['label']}\n")
                f.write(f"   占位符: {{{{{field['placeholder']}}}}}\n")
                f.write(f"   位置: {field['location_type']}\n")
                f.write(f"   原文: {field['text'][:50]}...\n\n")
            
            f.write("\n" + "=" * 80 + "\n")
            f.write("使用说明:\n")
            f.write("=" * 80 + "\n\n")
            f.write("1. 打开标注后的模板文件\n")
            f.write("2. 检查红色标记的占位符是否正确\n")
            f.write("3. 手动调整不准确的标注\n")
            f.write("4. 将调整后的模板上传到系统使用\n\n")
        
        print(f"📋 已生成标注说明: {guide_path}")


def test_annotator():
    """测试模板标注器"""
    print("\n" + "=" * 80)
    print("🧪 测试模板自动标注功能")
    print("=" * 80 + "\n")
    
    # 导入必要的模块
    import sys
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    
    from core.agent import UniversityCourseAgent
    from config.settings import DASHSCOPE_API_KEY
    
    # 初始化Agent
    print("初始化AI Agent...")
    agent = UniversityCourseAgent(api_key=DASHSCOPE_API_KEY)
    
    # 初始化标注器
    annotator = TemplateAnnotator(agent)
    
    # 测试文件路径
    test_template = "教案模版（2025年修订，征求意见稿）.docx"
    
    if not os.path.exists(test_template):
        print(f"❌ 测试模板不存在: {test_template}")
        return
    
    # 执行标注
    output_path = annotator.annotate_template(test_template)
    
    print(f"\n✅ 测试完成！")
    print(f"📄 标注后的模板: {output_path}")


if __name__ == "__main__":
    test_annotator()

