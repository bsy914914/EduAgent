"""
高级生成器 - Advanced Generator
使用AI自动填充教案模板的所有占位符
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import asyncio


class AdvancedLessonGenerator:
    """高级教案生成器"""
    
    # 占位符中文说明映射
    PLACEHOLDER_DESCRIPTIONS = {
        # 基本信息
        'course_name': '课程名称',
        'chapter_section': '授课章节',
        'class_name': '授课班级',
        'teacher_name': '授课教师',
        
        # 教学目标
        'ideological_goals': '思政育人目标（结合课程思政，培养学生的价值观、职业道德等）',
        'knowledge_goals': '知识目标（学生应掌握的知识点，通常3-5条）',
        'ability_goals': '能力目标（学生应获得的技能和能力，通常3-5条）',
        
        # 教学重难点
        'teaching_focus': '教学重点（本节课的核心知识点）',
        'teaching_difficulty': '教学难点（学生理解的难点）',
        'focus_solutions': '重点解决措施（如何突出重点的教学策略）',
        'difficulty_solutions': '难点解决措施（如何突破难点的教学策略）',
        
        # 教学方法与资源
        'teaching_methods': '教学方法（如讲授法、案例法、讨论法、演示法等）',
        'learning_methods': '学法指导（引导学生如何学习，如自主学习、合作学习等）',
        'teaching_resources': '教学资源（PPT、视频、案例、教材等）',
        'ideological_elements': '思政元素融入点（在哪些环节融入思政内容）',
        
        # 教学过程 - 预习环节
        'preview_content': '预习内容（学生课前需要预习的内容）',
        'preview_intention': '预习设计意图（为什么设计这样的预习任务）',
        'preview_teacher': '预习教师活动（教师在预习环节的指导工作）',
        'preview_student': '预习学生活动（学生在预习环节的具体任务）',
        
        # 教学过程 - 导入环节
        'introduction_content': '导入内容（如何引入新课，可以是案例、问题、故事等）',
        'introduction_intention': '导入设计意图（为什么这样导入）',
        'introduction_teacher': '导入教师活动（教师的具体导入行为）',
        'introduction_student': '导入学生活动（学生在导入环节的参与方式）',
        
        # 教学过程 - 新课讲授
        'teaching_content': '讲授内容（核心知识点的详细讲解，分点阐述）',
        'teaching_intention': '讲授设计意图（为什么这样安排讲授顺序和内容）',
        'teaching_teacher': '讲授教师活动（教师的讲解、演示、提问等活动）',
        'teaching_student': '讲授学生活动（学生的听讲、记录、思考、回答等活动）',
        
        # 教学过程 - 自主学习
        'self_learning_content': '自主学习内容（学生独立探索的内容）',
        'self_learning_intention': '自主学习设计意图（培养学生的自主学习能力）',
        'self_learning_teacher': '自主学习教师活动（教师的引导、观察、指导）',
        'self_learning_student': '自主学习学生活动（学生的阅读、思考、探索）',
        
        # 教学过程 - 课堂练习
        'practice_content': '练习内容（具体的练习题目或任务）',
        'practice_intention': '练习设计意图（巩固哪些知识点，训练哪些能力）',
        'practice_teacher': '练习教师活动（教师的巡视、指导、点评）',
        'practice_student': '练习学生活动（学生的独立练习、小组讨论）',
        
        # 教学过程 - 成果展示
        'presentation_content': '展示内容（学生展示的成果形式）',
        'presentation_intention': '展示设计意图（培养学生的表达能力和自信心）',
        'presentation_teacher': '展示教师活动（教师的组织、点评、总结）',
        'presentation_student': '展示学生活动（学生的展示、互评、反思）',
        
        # 教学过程 - 拓展延伸
        'extension_content': '拓展内容（课外延伸的知识或应用）',
        'extension_intention': '拓展设计意图（拓宽学生视野，激发学习兴趣）',
        'extension_teacher': '拓展教师活动（教师的推荐、引导）',
        'extension_student': '拓展学生活动（学生的课后探索、实践）',
        
        # 教学过程 - 课堂小结
        'evaluation_content': '小结内容（本节课的知识点梳理）',
        'evaluation_intention': '小结设计意图（帮助学生构建知识体系）',
        'evaluation_teacher': '小结教师活动（教师的总结、提炼、升华）',
        'evaluation_student': '小结学生活动（学生的回顾、归纳、反思）',
        
        # 教学过程 - 作业布置
        'homework_content': '作业内容（课后作业的具体要求）',
        'homework_intention': '作业设计意图（巩固知识、延伸学习）',
        'homework_teacher': '作业教师活动（教师的布置、要求说明）',
        'homework_student': '作业学生活动（学生的完成、提交）',
    }
    
    def __init__(self, agent=None, progress_callback=None):
        """
        初始化高级生成器
        
        Args:
            agent: UniversityCourseAgent实例，用于调用AI生成内容
            progress_callback: 进度回调函数 callback(progress, status, log_message)
        """
        self.agent = agent
        self.template_path = None
        self.progress_callback = progress_callback
        self.placeholders = []
    
    def _log_progress(self, progress, status, message):
        """记录进度"""
        print(f"[{progress}%] {message}")
        if self.progress_callback:
            self.progress_callback(progress, status, message)
        
    def analyze_template(self, template_path: str) -> List[str]:
        """
        分析模板文件，提取所有占位符
        
        Args:
            template_path: 模板文件路径
            
        Returns:
            占位符列表
        """
        print(f"📄 正在分析模板: {template_path}")
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        self.template_path = template_path
        doc = Document(template_path)
        placeholders = set()
        
        # 扫描段落
        for para in doc.paragraphs:
            matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
            placeholders.update(matches)
        
        # 扫描表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
                        placeholders.update(matches)
        
        self.placeholders = sorted(list(placeholders))
        
        print(f"✅ 发现 {len(self.placeholders)} 个占位符")
        
        return self.placeholders
    
    async def generate_content_for_placeholder(
        self, 
        placeholder: str, 
        topic: str,
        context: Dict[str, str] = None
    ) -> str:
        """
        为单个占位符生成内容
        
        Args:
            placeholder: 占位符名称
            topic: 教案主题
            context: 已生成的其他内容（用于保持一致性）
            
        Returns:
            生成的内容
        """
        description = self.PLACEHOLDER_DESCRIPTIONS.get(placeholder, placeholder)
        
        # 根据字段类型确定长度要求
        # 超短字段：姓名、班级等基本信息
        very_short_fields = ['teacher_name', 'class_name', 'course_name', 'chapter_section']
        # 短字段：方法、资源等
        short_fields = ['teaching_methods', 'learning_methods', 'teaching_resources', 
                       'teaching_focus', 'teaching_difficulty', 'ideological_elements']
        # 中等字段：目标、措施等
        medium_fields = ['knowledge_goals', 'ability_goals', 'ideological_goals',
                        'focus_solutions', 'difficulty_solutions']
        
        # 设置长度要求
        if placeholder in very_short_fields:
            if placeholder == 'chapter_section':
                length_req = "直接生成章节名称，例如：'第三章 数据结构' 或 '2.3 函数式编程'，不超过15字，不要用列表"
            elif placeholder == 'course_name':
                length_req = "直接生成课程名称，例如：'Python程序设计' 或 '数据库原理'，不超过10字，不要用列表，不要加编号"
            else:
                length_req = "直接生成具体内容，不超过10个字，不要解释，不要用列表"
        elif placeholder in short_fields:
            length_req = "用1-2句话简要说明，总共不超过30字"
        elif placeholder in medium_fields:
            length_req = "如果是列表，列出3-4条，每条不超过25字；如果是段落，不超过80字"
        else:
            length_req = "如果是列表，列出2-3条，每条不超过30字；如果是段落，控制在60-100字"
        
        # 构建提示词
        # 对于简短字段，强调不要用列表
        list_instruction = ""
        if placeholder in very_short_fields:
            list_instruction = "5. 不要使用列表格式，直接输出内容本身"
        else:
            list_instruction = "5. 如需列表，用数字编号（1. 2. 3.）"
        
        prompt = f"""请为教案主题「{topic}」生成以下内容：

字段名称：{placeholder}
字段说明：{description}

要求：
1. {length_req}
2. 内容要专业、简洁、具体
3. 符合高等教育教学规范
4. 与主题紧密相关
{list_instruction}

"""
        
        # 如果有上下文，添加到提示词中以保持一致性
        if context:
            if 'course_name' in context:
                prompt += f"\n课程名称：{context['course_name']}"
            if 'chapter_section' in context:
                prompt += f"\n授课章节：{context['chapter_section']}"
        
        prompt += f"\n\n请直接生成「{description}」的内容，不要包含字段名称："
        
        # 使用AI生成内容
        if self.agent:
            try:
                # 使用教案生成的LLM（速度较快）
                from langchain.schema import HumanMessage
                response = await self.agent.llm_lesson.ainvoke([HumanMessage(content=prompt)])
                content = response.content.strip()
                return content
            except Exception as e:
                print(f"⚠️  生成 {placeholder} 时出错: {e}")
                return f"[待填充: {description}]"
        else:
            return f"[待填充: {description}]"
    
    async def generate_all_content(self, topic: str) -> Dict[str, str]:
        """
        为所有占位符生成内容
        
        Args:
            topic: 教案主题
            
        Returns:
            占位符到内容的映射字典
        """
        print(f"\n🤖 开始为主题「{topic}」生成内容...")
        print(f"📝 需要生成 {len(self.placeholders)} 个字段的内容\n")
        
        self._log_progress(30, 'generating', f'🤖 开始生成内容，共需生成 {len(self.placeholders)} 个字段')
        
        content_dict = {}
        
        # 按优先级分组
        # 第一组：基本信息（先生成，其他内容会参考这些）
        basic_fields = ['course_name', 'chapter_section', 'class_name', 'teacher_name']
        
        # 第二组：目标和重难点
        goals_fields = ['ideological_goals', 'knowledge_goals', 'ability_goals', 
                       'teaching_focus', 'teaching_difficulty', 'focus_solutions', 'difficulty_solutions']
        
        # 第三组：方法和资源
        method_fields = ['teaching_methods', 'learning_methods', 'teaching_resources', 'ideological_elements']
        
        # 第四组：教学过程
        process_fields = [p for p in self.placeholders 
                         if p not in basic_fields + goals_fields + method_fields]
        
        # 按组顺序生成
        all_groups = [
            ("基本信息", basic_fields),
            ("教学目标与重难点", goals_fields),
            ("教学方法与资源", method_fields),
            ("教学过程", process_fields)
        ]
        
        total_fields = len(self.placeholders)
        completed_fields = 0
        
        for group_name, fields in all_groups:
            print(f"📌 生成 {group_name}...")
            self._log_progress(
                30 + int((completed_fields / total_fields) * 55),
                'generating',
                f'📌 正在生成{group_name}...'
            )
            
            for field in fields:
                if field not in self.placeholders:
                    continue
                
                field_desc = self.PLACEHOLDER_DESCRIPTIONS.get(field, field)
                print(f"   • 生成 {field}...", end=" ")
                
                content = await self.generate_content_for_placeholder(
                    field, topic, content_dict
                )
                
                content_dict[field] = content
                completed_fields += 1
                
                progress = 30 + int((completed_fields / total_fields) * 55)
                print(f"✓ ({len(content)} 字符)")
                self._log_progress(
                    progress,
                    'generating',
                    f'✓ 已生成「{field_desc}」({completed_fields}/{total_fields})'
                )
            
            print()
        
        print(f"✅ 所有内容生成完成！")
        self._log_progress(85, 'filling', '✅ 所有内容生成完成！开始填充模板...')
        
        return content_dict
    
    def fill_template(
        self, 
        content_dict: Dict[str, str], 
        output_path: str
    ) -> str:
        """
        用生成的内容填充模板
        
        Args:
            content_dict: 占位符到内容的映射
            output_path: 输出文件路径
            
        Returns:
            输出文件的绝对路径
        """
        print(f"\n📝 正在填充模板...")
        print(f"   模板路径: {self.template_path}")
        print(f"   输出路径: {output_path}")
        print(f"   待替换项: {len(content_dict)} 个")
        
        # 加载模板
        doc = Document(self.template_path)
        
        replacements_made = 0
        
        # 替换段落中的占位符
        for para in doc.paragraphs:
            original_text = para.text
            new_text = original_text
            
            # 对每个占位符进行替换
            for placeholder, content in content_dict.items():
                pattern = r'\{\{' + re.escape(placeholder) + r'\}\}'
                if re.search(pattern, new_text):
                    new_text = re.sub(pattern, content, new_text)
                    replacements_made += 1
                    print(f"   ✓ 替换段落中的 {{{{{placeholder}}}}}")
            
            # 如果文本有变化，清空段落并重新写入
            if new_text != original_text:
                # 保存第一个run的格式（字体、大小等，但不包括颜色）
                first_run = para.runs[0] if para.runs else None
                
                # 清空所有runs
                for run in para.runs:
                    run.text = ''
                
                # 写入新文本
                if first_run is not None:
                    first_run.text = new_text
                    # 明确设置为黑色字体
                    if first_run.font.color.rgb is not None:
                        first_run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    new_run = para.add_run(new_text)
                    # 设置为黑色字体
                    new_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        new_text = original_text
                        
                        # 对每个占位符进行替换
                        for placeholder, content in content_dict.items():
                            pattern = r'\{\{' + re.escape(placeholder) + r'\}\}'
                            if re.search(pattern, new_text):
                                new_text = re.sub(pattern, content, new_text)
                                replacements_made += 1
                                print(f"   ✓ 替换表格中的 {{{{{placeholder}}}}}")
                        
                        # 如果文本有变化，清空段落并重新写入
                        if new_text != original_text:
                            # 保存第一个run的格式（字体、大小等，但不包括颜色）
                            first_run = para.runs[0] if para.runs else None
                            
                            # 清空所有runs
                            for run in para.runs:
                                run.text = ''
                            
                            # 写入新文本
                            if first_run is not None:
                                first_run.text = new_text
                                # 明确设置为黑色字体
                                if first_run.font.color.rgb is not None:
                                    first_run.font.color.rgb = RGBColor(0, 0, 0)
                            else:
                                new_run = para.add_run(new_text)
                                # 设置为黑色字体
                                new_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 保存文件
        doc.save(output_path)
        
        print(f"✅ 模板填充完成！")
        print(f"   替换了 {replacements_made} 个占位符")
        print(f"   输出文件: {output_path}")
        
        return os.path.abspath(output_path)
    
    async def generate(
        self, 
        topic: str, 
        template_path: str = None,
        output_dir: str = None
    ) -> Tuple[bool, str]:
        """
        完整的生成流程
        
        Args:
            topic: 教案主题
            template_path: 模板文件路径（可选，默认使用标注教案.docx）
            output_dir: 输出目录（可选，默认为interface/exports）
            
        Returns:
            (成功标志, 输出文件路径)
        """
        try:
            # 确定模板路径
            if template_path is None:
                # 默认使用项目根目录的标注教案.docx
                project_root = Path(__file__).parent.parent
                template_path = project_root / "标注教案.docx"
            
            template_path = str(template_path)
            
            # 确定输出目录
            if output_dir is None:
                project_root = Path(__file__).parent.parent
                output_dir = project_root / "interface" / "exports"
            
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # 生成输出文件名
            safe_topic = re.sub(r'[^\w\s-]', '', topic).strip()
            safe_topic = re.sub(r'[-\s]+', '_', safe_topic)
            output_filename = f"{safe_topic}_教案_已完成.docx"
            output_path = output_dir / output_filename
            
            print("=" * 80)
            print("🚀 @高级生成 - Advanced Generation")
            print("=" * 80)
            print(f"📚 主题: {topic}")
            print(f"📄 模板: {template_path}")
            print(f"💾 输出: {output_path}")
            print("=" * 80)
            print()
            
            self._log_progress(10, 'uploading', f'📤 开始分析模板: {os.path.basename(template_path)}')
            
            # 步骤1：分析模板
            self.analyze_template(template_path)
            self._log_progress(20, 'analyzing', f'🔍 模板分析完成，发现 {len(self.placeholders)} 个占位符')
            
            # 步骤2：生成所有内容
            content_dict = await self.generate_all_content(topic)
            
            # 步骤3：填充模板
            self._log_progress(85, 'filling', '⚙️ 正在将内容填充到模板...')
            final_path = self.fill_template(content_dict, str(output_path))
            
            print()
            print("=" * 80)
            print("🎉 教案生成成功！")
            print("=" * 80)
            print(f"📁 文件位置: {final_path}")
            print(f"📊 生成字段: {len(content_dict)} 个")
            print(f"📝 总字数: {sum(len(v) for v in content_dict.values())} 字")
            print("=" * 80)
            
            self._log_progress(100, 'completed', f'🎉 教案生成成功！文件已保存')
            
            return True, final_path
            
        except Exception as e:
            print(f"\n❌ 生成失败: {e}")
            import traceback
            traceback.print_exc()
            return False, str(e)


async def advanced_generate(topic: str, agent=None) -> Tuple[bool, str]:
    """
    @高级生成 命令的快捷函数
    
    Args:
        topic: 教案主题
        agent: AI代理实例
        
    Returns:
        (成功标志, 输出文件路径或错误信息)
    """
    generator = AdvancedLessonGenerator(agent)
    return await generator.generate(topic)

